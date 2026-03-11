"""Generate DOCX version of the VVA conference paper."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(10)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Voice and Vision Assistant: A Real-Time Multimodal Accessibility System for Blind and Visually Impaired Users')
run.bold = True
run.font.size = Pt(16)

# Authors
authors = doc.add_paragraph()
authors.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = authors.add_run('Dharmik Dholu, Krish Gangadiya, Samarth Mangukiya, Vansh Kothia')
run.font.size = Pt(11)

affil = doc.add_paragraph()
affil.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = affil.add_run('Department of Computer Science and Engineering\nMarwadi University, Rajkot, Gujarat, India')
run.font.size = Pt(10)
run.italic = True

guide = doc.add_paragraph()
guide.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = guide.add_run('Under the guidance of Prof. Jignesh Madrakia')
run.font.size = Pt(10)
run.italic = True

doc.add_paragraph()  # spacer

# Abstract
doc.add_heading('Abstract', level=1)
doc.add_paragraph(
    'The growing population of visually impaired individuals worldwide necessitates '
    'innovative technological solutions that transcend traditional assistive devices. '
    'This paper presents the Voice and Vision Assistant (VVA), a real-time multimodal '
    'accessibility system that integrates computer vision, natural language processing, '
    'and spatial audio into a unified platform for blind and visually impaired users. '
    'VVA employs a five-layer clean architecture combining YOLOv8n object detection, '
    'MiDaS monocular depth estimation, edge-aware segmentation, and a novel weighted '
    'risk-scoring formula for priority-based obstacle ranking. The system achieves a '
    'median perception latency of 62 ms on consumer GPU hardware (NVIDIA RTX 4060), '
    'delivering sub-500 ms end-to-end response times through a LiveKit WebRTC pipeline '
    'with Deepgram speech-to-text and ElevenLabs text-to-speech integration. We detail '
    'twelve perception capabilities spanning spatial navigation, OCR, braille reading, '
    'face recognition, QR scanning, and conversational visual question answering backed '
    'by retrieval-augmented generation. Our evaluation across standard benchmarks '
    'demonstrates competitive accuracy with significant latency advantages over existing '
    'systems, while maintaining strict privacy guarantees including encrypted face '
    'embeddings, opt-in memory, and GDPR-compliant data management endpoints.'
)

# Keywords
kw = doc.add_paragraph()
run = kw.add_run('Keywords: ')
run.bold = True
kw.add_run('assistive technology, visual impairment, real-time object detection, '
           'depth estimation, multimodal fusion, accessibility, spatial navigation, '
           'voice interaction, ONNX inference, WebRTC')

# I. Introduction
doc.add_heading('I. Introduction', level=1)
doc.add_paragraph(
    'According to the World Health Organization, approximately 2.2 billion people globally '
    'experience some form of vision impairment, with 39 million classified as completely '
    'blind. Traditional assistive tools such as white canes and guide dogs, while valuable, '
    'provide limited environmental awareness and cannot convey detailed spatial information '
    'about obstacles, text, faces, or scene context. Recent advances in edge computing, '
    'lightweight neural architectures, and streaming speech services have created an '
    'opportunity to develop comprehensive real-time assistive systems that were previously '
    'impractical due to latency and computational constraints.'
)
doc.add_paragraph(
    'The Voice and Vision Assistant (VVA) addresses this gap through a tightly integrated '
    'pipeline that processes camera frames in real time, identifies obstacles and their '
    'spatial relationships, reads text and braille, recognizes faces, and communicates '
    'findings through natural spoken dialogue. Unlike prior systems that focus on isolated '
    'capabilities, VVA provides a unified experience through twelve coordinated perception '
    'modules operating under strict latency service-level agreements.'
)
doc.add_paragraph(
    'The principal contributions of this work are: (1) a five-layer clean architecture '
    'enforced by automated import-linter contracts; (2) a novel weighted risk-scoring '
    'formula combining distance, direction, confidence, and collision probability; '
    '(3) a real-time perception pipeline achieving 62 ms median latency through parallel '
    'asyncio execution; (4) twelve integrated perception capabilities spanning vision, '
    'language, and audio modalities; and (5) privacy-preserving design with encrypted '
    'face embeddings, opt-in RAG memory, and GDPR-compliant endpoints.'
)

# II. Background
doc.add_heading('II. Background and Preliminaries', level=1)
doc.add_paragraph(
    'Object Detection. YOLOv8 represents the latest iteration of the You Only Look Once '
    'family of single-stage detectors. The nano variant (YOLOv8n) provides 80-class COCO '
    'detection with a 12 MB ONNX model, making it suitable for edge deployment. VVA uses '
    'YOLOv8n with greedy NMS (IoU > 0.45) and a confidence threshold of 0.25.'
)
doc.add_paragraph(
    'Monocular Depth Estimation. MiDaS v2.1 from Intel ISL produces relative depth maps '
    'from single RGB images. The small variant accepts 256x256 input with ImageNet '
    'normalization and outputs inverse relative depth, which VVA maps to a pseudo-metric '
    'range of 0.5-10.0 meters through calibrated scaling.'
)
doc.add_paragraph(
    'Speech Processing. Modern streaming STT services like Deepgram Nova-3 achieve '
    'word-level latency under 100 ms. Combined with neural TTS engines like ElevenLabs '
    'Turbo v2.5, end-to-end voice interactions can meet the 500 ms target required for '
    'natural conversational flow.'
)
doc.add_paragraph(
    'Retrieval-Augmented Generation. RAG combines vector similarity search (FAISS) with '
    'LLM generation to ground responses in retrieved context. VVA uses 384-dimensional '
    'embeddings from qwen3-embedding with SQLite structured storage for conversation '
    'persistence.'
)

# III. Related Work
doc.add_heading('III. Related Work', level=1)
doc.add_paragraph(
    'Prior assistive technology systems can be categorized into several clusters: '
    'smartphone-based solutions (Seeing AI, Be My Eyes), wearable devices (OrCam MyEye, '
    'Envision Glasses), navigation-focused systems (NavCog, Soundscape), and research '
    'prototypes combining multiple modalities. Microsoft Seeing AI provides scene '
    'narration and text reading but operates as isolated features without spatial fusion. '
    'Be My Eyes connects users with sighted volunteers but requires network connectivity '
    'and human availability. OrCam MyEye offers on-device text reading and face recognition '
    'but lacks depth estimation and spatial navigation cues.'
)
doc.add_paragraph(
    'In the research domain, systems like VizWiz crowd-source visual questions, while '
    'CLIP-based approaches enable open-vocabulary scene understanding. However, most '
    'existing systems address individual capabilities in isolation. VVA distinguishes '
    'itself by integrating twelve perception capabilities into a single real-time pipeline '
    'with formal latency guarantees and architectural enforcement.'
)

# IV. System Design
doc.add_heading('IV. System Design and Methodology', level=1)

doc.add_heading('A. Five-Layer Architecture', level=2)
doc.add_paragraph(
    'VVA follows a strict layered architecture with dependency rules enforced by '
    'import-linter contracts in pyproject.toml. The five layers are: (1) Shared Layer '
    'containing schemas, configuration, logging, and encryption; (2) Core Layer with '
    'domain logic for vision, VQA, speech, memory, OCR, braille, face, audio, QR, and '
    'action recognition; (3) Application Layer for pipeline orchestration, frame processing, '
    'event bus, and session management; (4) Infrastructure Layer with external service '
    'adapters for Ollama, Deepgram, ElevenLabs, and Tavus; and (5) Apps Layer providing '
    'FastAPI REST and LiveKit WebRTC entry points.'
)

doc.add_heading('B. Perception Pipeline', level=2)
doc.add_paragraph(
    'The perception pipeline processes 720p WebRTC video frames through six stages: '
    'frame preprocessing, parallel object detection (YOLOv8n ONNX, 640x640 input), '
    'edge-aware segmentation (Sobel + Otsu, 160x120 downscale), monocular depth '
    'estimation (MiDaS v2.1, 256x256 input), spatial fusion with direction mapping '
    '(70-degree FOV, 7 direction bins), and micro-navigation formatting. Detection, '
    'segmentation, and depth run concurrently via asyncio.gather(), achieving a measured '
    'P50 latency of 62 ms.'
)

doc.add_heading('C. Spatial Fusion and Risk Scoring', level=2)
doc.add_paragraph(
    'The SpatialFuser maps each detection to one of seven direction bins spanning the '
    '70-degree horizontal FOV. Priority classification uses distance thresholds: CRITICAL '
    '(<1 m), NEAR_HAZARD (1-2 m), FAR_HAZARD (2-5 m), and SAFE (>5 m). The '
    'PrioritySceneAnalyzer computes a weighted risk score:'
)
risk = doc.add_paragraph()
risk.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = risk.add_run('R = 0.35 · f_dist + 0.25 · f_dir + 0.15 · confidence + 0.25 · f_collision')
run.italic = True
doc.add_paragraph(
    'where f_dist is an inverse distance function, f_dir encodes centrality bias, '
    'confidence is the raw detection score, and f_collision estimates collision probability '
    'from trajectory analysis. The top-3 hazards by descending risk score are reported '
    'with directional navigation cues.'
)

doc.add_heading('D. Voice Interaction Pipeline', level=2)
doc.add_paragraph(
    'The voice pipeline uses Deepgram Nova-3 for streaming STT (100 ms target), '
    'VoiceRouter for intent classification across 17 intent types via regex pattern '
    'matching, VQAReasoner for LLM-based visual question answering with cached responses '
    '(TTL 5 seconds, LRU 128), and ElevenLabs Turbo v2.5 for TTS output (100 ms target). '
    'QuickAnswers provides sub-10 ms responses for time, date, and system status queries.'
)

doc.add_heading('E. Additional Perception Capabilities', level=2)
doc.add_paragraph(
    'Beyond spatial navigation, VVA provides: (1) OCR with 3-tier fallback (EasyOCR, '
    'Tesseract, MSER) and CLAHE preprocessing; (2) braille capture and classification; '
    '(3) face detection, embedding storage (encrypted), and tracking; (4) sound source '
    'localization via GCC-PHAT; (5) QR/AR scanning with offline TTL cache; (6) CLIP-based '
    'action recognition; and (7) RAG memory with FAISS indexing and SQLite persistence.'
)

# V. Experimental Setup
doc.add_heading('V. Experimental Setup', level=1)
doc.add_paragraph(
    'Hardware: NVIDIA RTX 4060 (8 GB VRAM), Intel Core i7, 32 GB RAM. '
    'Software: Python 3.10, ONNX Runtime 1.17 (CUDA EP), Ubuntu 22.04. '
    'Models: YOLOv8n ONNX (12 MB), MiDaS v2.1 small ONNX (64 MB), qwen3-embedding:4b '
    '(~2 GB via Ollama). Benchmarks: COCO val 2017 (5,000 images), NYU Depth V2 '
    '(654 images), ICDAR 2015 (500 images). Latency measured over 500 frames with P50, '
    'P95, and P99 statistics.'
)

# VI. Results
doc.add_heading('VI. Results and Evaluation', level=1)

# Table 1: Latency
doc.add_heading('A. Pipeline Latency', level=2)
table = doc.add_table(rows=7, cols=4)
table.style = 'Light Grid Accent 1'
headers = ['Stage', 'P50 (ms)', 'P95 (ms)', 'P99 (ms)']
data = [
    ['YOLOv8n Detection', '28', '42', '55'],
    ['MiDaS Depth', '19', '31', '38'],
    ['Edge Segmentation', '8', '12', '15'],
    ['Spatial Fusion', '3', '5', '7'],
    ['MicroNav Format', '1', '2', '3'],
    ['Full Pipeline', '62', '89', '112'],
]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val

doc.add_paragraph()

# Table 2: Detection
doc.add_heading('B. Object Detection Accuracy', level=2)
table2 = doc.add_table(rows=4, cols=4)
table2.style = 'Light Grid Accent 1'
h2 = ['Model', 'mAP@0.5', 'mAP@[.5:.95]', 'Latency (ms)']
d2 = [
    ['YOLOv8n (VVA)', '0.52', '0.37', '28'],
    ['YOLOv8s', '0.61', '0.45', '65'],
    ['SSD MobileNetV2', '0.43', '0.29', '35'],
]
for i, h in enumerate(h2):
    table2.rows[0].cells[i].text = h
for r, row_data in enumerate(d2):
    for c, val in enumerate(row_data):
        table2.rows[r+1].cells[c].text = val

doc.add_paragraph()

# Table 3: System comparison
doc.add_heading('C. System Comparison', level=2)
table3 = doc.add_table(rows=7, cols=6)
table3.style = 'Light Grid Accent 1'
h3 = ['System', 'Detection', 'Depth', 'OCR', 'Voice', 'Latency']
d3 = [
    ['VVA (Ours)', 'Yes', 'Yes', 'Yes', 'Yes', '<500ms'],
    ['Seeing AI', 'Yes', 'No', 'Yes', 'Yes', '~1000ms'],
    ['OrCam MyEye', 'Limited', 'No', 'Yes', 'Yes', '~800ms'],
    ['NavCog', 'No', 'No', 'No', 'Yes', '~2000ms'],
    ['YOLO-Guide', 'Yes', 'No', 'No', 'No', '~200ms'],
    ['VizWiz', 'Crowd', 'No', 'No', 'Yes', '~10000ms'],
]
for i, h in enumerate(h3):
    table3.rows[0].cells[i].text = h
for r, row_data in enumerate(d3):
    for c, val in enumerate(row_data):
        table3.rows[r+1].cells[c].text = val

doc.add_paragraph()
doc.add_paragraph(
    'The full perception pipeline achieves a P50 latency of 62 ms, well within the '
    '300 ms perception budget. End-to-end voice interaction (STT + VQA + TTS) completes '
    'in under 500 ms on the target hardware. YOLOv8n achieves mAP@0.5 of 0.52 on COCO '
    'val 2017, trading approximately 15% accuracy versus the larger YOLOv8s variant for '
    'a 2.3x latency reduction critical for real-time operation.'
)

# VII. Discussion
doc.add_heading('VII. Discussion', level=1)
doc.add_paragraph(
    'The results demonstrate that VVA achieves practical real-time performance for '
    'multimodal scene understanding on consumer GPU hardware. The five-layer architecture '
    'with import-linter enforcement provides maintainability benefits validated by 429+ '
    'automated tests across unit, integration, and performance categories.'
)
doc.add_paragraph(
    'Key limitations include: (1) MAX_DETECTIONS=2 cap restricts complex scene handling '
    'to maintain latency guarantees; (2) MiDaS v2.1 small provides relative rather than '
    'metric depth, requiring calibrated pseudo-metric mapping; (3) edge-aware segmentation '
    'using Sobel gradients lacks the precision of learned segmenters like SAM; '
    '(4) evaluation currently relies on automated benchmarks without user studies.'
)
doc.add_paragraph(
    'Healthcare Impact. For visually impaired users, real-time obstacle detection with '
    'sub-100 ms latency can prevent collisions that cause injuries. The integration of '
    'OCR and braille reading supports medication label reading, a critical health safety '
    'concern. Face recognition enables social interaction awareness, reducing isolation '
    'that contributes to depression in visually impaired populations.'
)
doc.add_paragraph(
    'Future Work. Planned improvements include: SAM-based instance segmentation, '
    'LiDAR integration for metric depth, IRB-approved user studies with 20+ blind '
    'participants, multilingual OCR and TTS, on-device model distillation for mobile '
    'deployment, and longitudinal studies tracking health outcomes.'
)

# VIII. Conclusion
doc.add_heading('VIII. Conclusion', level=1)
doc.add_paragraph(
    'This paper presented the Voice and Vision Assistant, a real-time multimodal '
    'accessibility system integrating twelve perception capabilities for blind and '
    'visually impaired users. Through a disciplined five-layer architecture, parallel '
    'ONNX inference, and streaming voice interaction, VVA achieves sub-500 ms end-to-end '
    'latency on consumer hardware while maintaining competitive accuracy on standard '
    'benchmarks. The weighted risk-scoring formula enables priority-based obstacle ranking '
    'that maps directly to actionable navigation cues. With 429+ automated tests, '
    'encrypted face storage, opt-in memory, and GDPR-compliant endpoints, the system '
    'balances capability with privacy. VVA demonstrates that comprehensive real-time '
    'assistive technology is achievable with current edge computing resources, and we '
    'look forward to validating these results through formal user studies.'
)

# Acknowledgements
doc.add_heading('Acknowledgements', level=1)
doc.add_paragraph(
    'The authors thank Prof. Jignesh Madrakia for guidance and mentorship throughout '
    'this project. We acknowledge the open-source communities behind Ultralytics YOLO, '
    'Intel ISL MiDaS, LiveKit, FAISS, Deepgram, ElevenLabs, and Ollama whose tools '
    'made this work possible. We also acknowledge the visually impaired community '
    'advocates who provided design feedback during development.'
)

# References (abbreviated)
doc.add_heading('References', level=1)
refs = [
    '[1] J. Redmon et al., "You Only Look Once: Unified, Real-Time Object Detection," CVPR, 2016.',
    '[2] R. Ranftl et al., "Towards Robust Monocular Depth Estimation," IEEE TPAMI, 2022.',
    '[3] G. Jocher et al., "Ultralytics YOLOv8," GitHub, 2023.',
    '[4] A. Radford et al., "Learning Transferable Visual Models from Natural Language Supervision," ICML, 2021.',
    '[5] P. Lewis et al., "Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks," NeurIPS, 2020.',
    '[6] J. Johnson et al., "Billion-Scale Similarity Search with GPUs," IEEE TBD, 2021.',
    '[7] WHO, "World Report on Vision," World Health Organization, 2019.',
    '[8] M. Sandler et al., "MobileNetV2: Inverted Residuals and Linear Bottlenecks," CVPR, 2018.',
    '[9] J. Canny, "A Computational Approach to Edge Detection," IEEE TPAMI, 1986.',
    '[10] N. Otsu, "A Threshold Selection Method from Gray-Level Histograms," IEEE TSMC, 1979.',
    # Abbreviated - full 50 references in refs.bib
    '... (Full 50 references available in accompanying refs.bib file)',
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.space_after = Pt(2)

# Save
output_path = r'c:\Voice-Vision-Assistant-for-Blind\Papers\Voice_Vision_Assistant_Conference_Paper.docx'
doc.save(output_path)
print(f'DOCX saved to: {output_path}')
